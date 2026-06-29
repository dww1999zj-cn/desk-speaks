#!/usr/bin/env python3
"""将 docs/PROJECT.md 导出为 Word (.docx) 和 PDF（需 reportlab）。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
MD_PATH = DOCS / "PROJECT.md"
DOCX_PATH = DOCS / "PROJECT.docx"
PDF_PATH = DOCS / "PROJECT.pdf"


def parse_markdown_blocks(text: str) -> list[tuple[str, str]]:
    """简单 Markdown 分块：heading / paragraph / table / code."""
    blocks: list[tuple[str, str]] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            i += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            blocks.append((f"h{min(level, 4)}", title))
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", "\n".join(table_lines)))
            continue

        if line.strip() == "---":
            blocks.append(("hr", ""))
            i += 1
            continue

        if line.strip():
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "```", "---")):
                para_lines.append(lines[i])
                i += 1
            blocks.append(("p", " ".join(para_lines)))
            continue

        i += 1

    return blocks


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def export_docx(blocks: list[tuple[str, str]]) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    for kind, content in blocks:
        if kind == "h1":
            doc.add_heading(strip_md_inline(content), level=0)
        elif kind == "h2":
            doc.add_heading(strip_md_inline(content), level=1)
        elif kind == "h3":
            doc.add_heading(strip_md_inline(content), level=2)
        elif kind == "h4":
            doc.add_heading(strip_md_inline(content), level=3)
        elif kind == "hr":
            doc.add_paragraph("—" * 40)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif kind == "table":
            rows = [r.strip() for r in content.splitlines() if r.strip()]
            if len(rows) < 2:
                continue
            header = [c.strip() for c in rows[0].strip("|").split("|")]
            data_rows = rows[2:]  # skip separator
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, cell in enumerate(header):
                table.rows[0].cells[j].text = strip_md_inline(cell)
            for row in data_rows:
                cells = [c.strip() for c in row.strip("|").split("|")]
                if len(cells) != len(header):
                    continue
                tr = table.add_row()
                for j, cell in enumerate(cells):
                    tr.cells[j].text = strip_md_inline(cell)
            doc.add_paragraph("")
        elif kind == "p":
            text = strip_md_inline(content)
            if text.startswith("*") and text.endswith("*"):
                p = doc.add_paragraph()
                run = p.add_run(text.strip("*"))
                run.italic = True
            else:
                doc.add_paragraph(text)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


def export_pdf(blocks: list[tuple[str, str]]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "CNBody",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=16,
    )
    h1 = ParagraphStyle("CNH1", parent=base, fontSize=18, leading=24, spaceAfter=12)
    h2 = ParagraphStyle("CNH2", parent=base, fontSize=14, leading=20, spaceBefore=12, spaceAfter=8)
    h3 = ParagraphStyle("CNH3", parent=base, fontSize=12, leading=18, spaceBefore=8, spaceAfter=6)
    code_style = ParagraphStyle("CNCode", parent=base, fontSize=8, leading=12, backColor=colors.whitesmoke)

    story = []

    def esc(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    for kind, content in blocks:
        if kind == "h1":
            story.append(Paragraph(esc(strip_md_inline(content)), h1))
        elif kind == "h2":
            story.append(Paragraph(esc(strip_md_inline(content)), h2))
        elif kind == "h3":
            story.append(Paragraph(esc(strip_md_inline(content)), h3))
        elif kind == "h4":
            story.append(Paragraph(esc(strip_md_inline(content)), h3))
        elif kind == "hr":
            story.append(Spacer(1, 8))
        elif kind == "code":
            story.append(Preformatted(content, code_style))
            story.append(Spacer(1, 6))
        elif kind == "table":
            rows = [r.strip() for r in content.splitlines() if r.strip()]
            if len(rows) < 2:
                continue
            header = [strip_md_inline(c.strip()) for c in rows[0].strip("|").split("|")]
            data = []
            data.append(header)
            for row in rows[2:]:
                cells = [strip_md_inline(c.strip()) for c in row.strip("|").split("|")]
                if len(cells) == len(header):
                    data.append(cells)
            t = Table(data, repeatRows=1)
            t.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F3EEFF")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 10))
        elif kind == "p":
            story.append(Paragraph(esc(strip_md_inline(content)), base))
            story.append(Spacer(1, 4))

    doc.build(story)
    print(f"Wrote {PDF_PATH}")


def main() -> int:
    if not MD_PATH.exists():
        print(f"Missing {MD_PATH}", file=sys.stderr)
        return 1

    text = MD_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown_blocks(text)

    try:
        export_docx(blocks)
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        import subprocess

        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        export_docx(blocks)

    try:
        export_pdf(blocks)
    except ImportError:
        print("Installing reportlab...", file=sys.stderr)
        import subprocess

        subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab", "-q"])
        export_pdf(blocks)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
